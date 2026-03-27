import io
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Flowable, Paragraph, SimpleDocTemplate, Spacer

from backend.utils import MedicalCard


class DocxGenerator:
    def __init__(self, medical_card: MedicalCard, source_transcript: str = ""):
        self.medical_card = medical_card
        self.source_transcript = source_transcript
        self.document = Document()

    def build(self) -> bytes:
        self._add_generated_at()
        self._add_title()
        self._add_patient_section()
        self._add_complaints_section()
        self._add_current_medications_section()
        self._add_diagnosis_section()
        self._add_prescriptions_section()
        self._add_patient_summary_section()
        self._add_doctor_summary_section()
        self._add_source_transcript_section()
        self._add_missing_info_section()
        return self._to_bytes()

    def _add_generated_at(self) -> None:
        self.document.add_paragraph(self._formatted_datetime())

    def _add_title(self) -> None:
        self.document.add_heading("Медична картка", level=1)

    def _add_patient_section(self) -> None:
        patient = self.medical_card.patient
        self.document.add_paragraph(f"ПІБ пацієнта: {patient.full_name or ''}")
        self.document.add_paragraph(f"Вік: {patient.age if patient.age is not None else ''}")
        self.document.add_paragraph(f"Стать: {patient.gender or ''}")

    def _add_complaints_section(self) -> None:
        self._add_section_heading("Скарги")
        for complaint in self.medical_card.complaints:
            if complaint.duration:
                self.document.add_paragraph(f"- {complaint.description} ({complaint.duration})")
            else:
                self.document.add_paragraph(f"- {complaint.description}")

    def _add_current_medications_section(self) -> None:
        self._add_section_heading("Поточні препарати")
        for medication in self.medical_card.current_medications:
            medication_name = medication.name or "Не зазначено"
            if medication.duration:
                self.document.add_paragraph(f"- {medication_name} ({medication.duration})")
            else:
                self.document.add_paragraph(f"- {medication_name}")

    def _add_diagnosis_section(self) -> None:
        if not self.medical_card.diagnosis:
            return
        self._add_section_heading("Діагноз")
        self.document.add_paragraph(self.medical_card.diagnosis.preliminary or "")

    def _add_prescriptions_section(self) -> None:
        self._add_section_heading("Призначення")
        for prescription in self.medical_card.prescriptions:
            details = []
            if prescription.dosage:
                details.append(prescription.dosage)
            if prescription.duration:
                details.append(prescription.duration)
            if details:
                self.document.add_paragraph(f"- {prescription.name} ({', '.join(details)})")
            else:
                self.document.add_paragraph(f"- {prescription.name}")

    def _add_patient_summary_section(self) -> None:
        if not self.medical_card.patient_summary:
            return
        self._add_section_heading("Підсумок скарг пацієнта")
        self.document.add_paragraph(self.medical_card.patient_summary)

    def _add_doctor_summary_section(self) -> None:
        if not self.medical_card.doctor_summary:
            return
        self._add_section_heading("Підсумок слів лікаря")
        self.document.add_paragraph(self.medical_card.doctor_summary)

    def _add_source_transcript_section(self) -> None:
        self._add_section_heading("Транскрипт запису")
        transcript_turns = self._split_transcript_turns()
        if not transcript_turns:
            self.document.add_paragraph("")
            return

        for index, turn in enumerate(transcript_turns):
            self._add_docx_dialogue_turn(turn)
            if index < len(transcript_turns) - 1:
                # Empty line between dialogue turns for readability.
                self.document.add_paragraph("")

    def _add_missing_info_section(self) -> None:
        missing_items = self._collect_missing_info()
        if not missing_items:
            return
        self._add_section_heading("Відсутня інформація")
        for item in missing_items:
            self.document.add_paragraph(f"- {item}")

    def _collect_missing_info(self) -> list[str]:
        missing: list[str] = []
        patient = self.medical_card.patient
        if not patient.full_name:
            missing.append("Відсутня інформація про ПІБ пацієнта")
        if patient.age is None:
            missing.append("Відсутня інформація про вік")
        if not patient.gender:
            missing.append("Відсутня інформація про стать")
        if not self.medical_card.complaints:
            missing.append("Відсутня інформація про скарги")
        elif any(not complaint.duration for complaint in self.medical_card.complaints):
            missing.append("Відсутня інформація про тривалість скарг")
        if not self.medical_card.current_medications:
            missing.append("Відсутня інформація про поточні препарати")
        elif any(not medication.duration for medication in self.medical_card.current_medications):
            missing.append("Відсутня інформація про тривалість прийому поточних препаратів")
        if not self.medical_card.diagnosis or not self.medical_card.diagnosis.preliminary:
            missing.append("Відсутня інформація про діагноз")
        if not self.medical_card.prescriptions:
            missing.append("Відсутня інформація про призначення")
        else:
            if any(not prescription.dosage for prescription in self.medical_card.prescriptions):
                missing.append("Відсутня інформація про дозування призначень")
            if any(not prescription.duration for prescription in self.medical_card.prescriptions):
                missing.append("Відсутня інформація про тривалість призначень")
        if not self.medical_card.patient_summary:
            missing.append("Відсутня інформація про підсумок скарг пацієнта")
        if not self.medical_card.doctor_summary:
            missing.append("Відсутня інформація про підсумок слів лікаря")
        if not self.source_transcript:
            missing.append("Відсутня інформація про транскрипт запису")
        return missing

    def _to_bytes(self) -> bytes:
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _formatted_datetime(self) -> str:
        generated_at = datetime.now().replace(
            year=self.medical_card.document_date.year,
            month=self.medical_card.document_date.month,
            day=self.medical_card.document_date.day,
        )
        return generated_at.strftime("%H:%M %d.%m.%Y")

    def _split_transcript_turns(self) -> list[str]:
        if not self.source_transcript:
            return []
        normalized = self.source_transcript.replace("\r\n", "\n")
        turns = [turn.strip() for turn in normalized.split("\n\n")]
        return [turn for turn in turns if turn]

    def _add_section_heading(self, text: str) -> None:
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True

    def _add_docx_dialogue_turn(self, turn: str) -> None:
        paragraph = self.document.add_paragraph()
        speaker, message = self._split_speaker_turn(turn)
        if speaker in {"Лікар", "Пацієнт"}:
            speaker_run = paragraph.add_run(f"{speaker}:")
            speaker_run.bold = True
            if message:
                paragraph.add_run(f" {message}")
            return
        paragraph.add_run(turn)

    def _split_speaker_turn(self, turn: str) -> tuple[str, str]:
        speaker, separator, message = turn.partition(":")
        if not separator:
            return "", turn.strip()
        return speaker.strip(), message.strip()


class PdfGenerator:
    def __init__(self, medical_card: MedicalCard, source_transcript: str = ""):
        self.medical_card = medical_card
        self.source_transcript = source_transcript
        self.buffer = io.BytesIO()
        self.content: list[Flowable] = []
        self.document = SimpleDocTemplate(self.buffer, pagesize=A4)

        self.font_name = "DejaVu"
        self.font_bold_name = "DejaVu-Bold"
        font_dir = Path(__file__).resolve().parent / "fonts"
        self.font_path = font_dir / "DejaVuSans.ttf"
        project_bold_path = font_dir / "DejaVuSans-Bold.ttf"
        system_bold_path = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        self.font_bold_path = project_bold_path if project_bold_path.exists() else system_bold_path
        self.title_style: ParagraphStyle | None = None
        self.heading_style: ParagraphStyle | None = None
        self.normal_style: ParagraphStyle | None = None

    def build(self) -> bytes:
        self._setup_font()
        self._setup_styles()
        self._add_generated_at()
        self._add_title()
        self._add_patient_section()
        self._add_complaints_section()
        self._add_current_medications_section()
        self._add_diagnosis_section()
        self._add_prescriptions_section()
        self._add_patient_summary_section()
        self._add_doctor_summary_section()
        self._add_source_transcript_section()
        self._add_missing_info_section()
        return self._to_bytes()

    def _setup_font(self) -> None:
        pdfmetrics.registerFont(TTFont(self.font_name, str(self.font_path)))
        if self.font_bold_path.exists():
            pdfmetrics.registerFont(TTFont(self.font_bold_name, str(self.font_bold_path)))
            return
        self.font_bold_name = self.font_name

    def _setup_styles(self) -> None:
        self.title_style = ParagraphStyle(
            name="TitleStyle",
            fontName=self.font_bold_name,
            fontSize=16,
            spaceAfter=14,
            alignment=1,
        )
        self.heading_style = ParagraphStyle(
            name="HeadingStyle",
            fontName=self.font_bold_name,
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
        )
        self.normal_style = ParagraphStyle(
            name="NormalStyle",
            fontName=self.font_name,
            fontSize=10,
            spaceAfter=4,
        )

    def _add_generated_at(self) -> None:
        self.content.append(Paragraph(self._formatted_datetime(), self.normal_style))

    def _add_title(self) -> None:
        self.content.append(Paragraph("Медична картка", self.title_style))

    def _add_patient_section(self) -> None:
        patient = self.medical_card.patient
        self.content.append(Paragraph(f"ПІБ пацієнта: {patient.full_name or ''}", self.normal_style))
        self.content.append(Paragraph(f"Вік: {patient.age if patient.age is not None else ''}", self.normal_style))
        self.content.append(Paragraph(f"Стать: {patient.gender or ''}", self.normal_style))

    def _add_complaints_section(self) -> None:
        self.content.append(Paragraph("Скарги", self.heading_style))
        for complaint in self.medical_card.complaints:
            if complaint.duration:
                self.content.append(Paragraph(f"- {complaint.description} ({complaint.duration})", self.normal_style))
            else:
                self.content.append(Paragraph(f"- {complaint.description}", self.normal_style))

    def _add_current_medications_section(self) -> None:
        self.content.append(Paragraph("Поточні препарати", self.heading_style))
        for medication in self.medical_card.current_medications:
            medication_name = medication.name or "Не зазначено"
            if medication.duration:
                self.content.append(Paragraph(f"- {medication_name} ({medication.duration})", self.normal_style))
            else:
                self.content.append(Paragraph(f"- {medication_name}", self.normal_style))

    def _add_diagnosis_section(self) -> None:
        if not self.medical_card.diagnosis:
            return
        self.content.append(Paragraph("Діагноз", self.heading_style))
        self.content.append(Paragraph(self.medical_card.diagnosis.preliminary or "", self.normal_style))

    def _add_prescriptions_section(self) -> None:
        self.content.append(Paragraph("Призначення", self.heading_style))
        for prescription in self.medical_card.prescriptions:
            details = []
            if prescription.dosage:
                details.append(prescription.dosage)
            if prescription.duration:
                details.append(prescription.duration)
            if details:
                self.content.append(Paragraph(f"- {prescription.name} ({', '.join(details)})", self.normal_style))
            else:
                self.content.append(Paragraph(f"- {prescription.name}", self.normal_style))

    def _add_patient_summary_section(self) -> None:
        if not self.medical_card.patient_summary:
            return
        self.content.append(Paragraph("Підсумок скарг пацієнта", self.heading_style))
        self.content.append(Paragraph(self.medical_card.patient_summary, self.normal_style))

    def _add_doctor_summary_section(self) -> None:
        if not self.medical_card.doctor_summary:
            return
        self.content.append(Paragraph("Підсумок слів лікаря", self.heading_style))
        self.content.append(Paragraph(self.medical_card.doctor_summary, self.normal_style))

    def _add_source_transcript_section(self) -> None:
        self.content.append(Paragraph("Транскрипт запису", self.heading_style))
        transcript_turns = self._split_transcript_turns()
        if not transcript_turns:
            self.content.append(Paragraph("", self.normal_style))
            return

        for index, turn in enumerate(transcript_turns):
            self.content.append(Paragraph(self._format_pdf_dialogue_turn(turn), self.normal_style))
            if index < len(transcript_turns) - 1:
                # Empty line between dialogue turns for readability.
                self.content.append(Spacer(1, 6))

    def _add_missing_info_section(self) -> None:
        missing_items = self._collect_missing_info()
        if not missing_items:
            return
        self.content.append(Paragraph("Відсутня інформація", self.heading_style))
        for item in missing_items:
            self.content.append(Paragraph(f"- {item}", self.normal_style))

    def _collect_missing_info(self) -> list[str]:
        missing: list[str] = []
        patient = self.medical_card.patient
        if not patient.full_name:
            missing.append("Відсутня інформація про ПІБ пацієнта")
        if patient.age is None:
            missing.append("Відсутня інформація про вік")
        if not patient.gender:
            missing.append("Відсутня інформація про стать")
        if not self.medical_card.complaints:
            missing.append("Відсутня інформація про скарги")
        elif any(not complaint.duration for complaint in self.medical_card.complaints):
            missing.append("Відсутня інформація про тривалість скарг")
        if not self.medical_card.current_medications:
            missing.append("Відсутня інформація про поточні препарати")
        elif any(not medication.duration for medication in self.medical_card.current_medications):
            missing.append("Відсутня інформація про тривалість прийому поточних препаратів")
        if not self.medical_card.diagnosis or not self.medical_card.diagnosis.preliminary:
            missing.append("Відсутня інформація про діагноз")
        if not self.medical_card.prescriptions:
            missing.append("Відсутня інформація про призначення")
        else:
            if any(not prescription.dosage for prescription in self.medical_card.prescriptions):
                missing.append("Відсутня інформація про дозування призначень")
            if any(not prescription.duration for prescription in self.medical_card.prescriptions):
                missing.append("Відсутня інформація про тривалість призначень")
        if not self.medical_card.patient_summary:
            missing.append("Відсутня інформація про підсумок скарг пацієнта")
        if not self.medical_card.doctor_summary:
            missing.append("Відсутня інформація про підсумок слів лікаря")
        if not self.source_transcript:
            missing.append("Відсутня інформація про транскрипт запису")
        return missing

    def _to_bytes(self) -> bytes:
        self.document.build(self.content)
        self.buffer.seek(0)
        return self.buffer.read()

    def _formatted_datetime(self) -> str:
        generated_at = datetime.now().replace(
            year=self.medical_card.document_date.year,
            month=self.medical_card.document_date.month,
            day=self.medical_card.document_date.day,
        )
        return generated_at.strftime("%H:%M %d.%m.%Y")

    def _split_transcript_turns(self) -> list[str]:
        if not self.source_transcript:
            return []
        normalized = self.source_transcript.replace("\r\n", "\n")
        turns = [turn.strip() for turn in normalized.split("\n\n")]
        return [turn for turn in turns if turn]

    def _format_pdf_dialogue_turn(self, turn: str) -> str:
        speaker, message = self._split_speaker_turn(turn)
        if speaker in {"Лікар", "Пацієнт"}:
            safe_speaker = escape(speaker)
            safe_message = escape(message)
            return f"<b>{safe_speaker}:</b> {safe_message}".strip()
        return escape(turn)

    def _split_speaker_turn(self, turn: str) -> tuple[str, str]:
        speaker, separator, message = turn.partition(":")
        if not separator:
            return "", turn.strip()
        return speaker.strip(), message.strip()


def generate_docx(medical_card: MedicalCard, source_transcript: str = "") -> bytes:
    return DocxGenerator(medical_card, source_transcript).build()


def generate_pdf(medical_card: MedicalCard, source_transcript: str = "") -> bytes:
    return PdfGenerator(medical_card, source_transcript).build()
