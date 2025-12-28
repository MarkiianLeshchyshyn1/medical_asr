from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
from utils import MedicalCard


def generate_docx(medical_card: MedicalCard) -> bytes:
    doc = Document()

    doc.add_heading("Medical Record", level=1)

    # Patient
    doc.add_paragraph(f"Patient name: {medical_card.patient.full_name}")
    doc.add_paragraph(f"Age: {medical_card.patient.age}")
    doc.add_paragraph(f"Gender: {medical_card.patient.gender}")

    doc.add_heading("Complaints", level=2)
    for c in medical_card.complaints:
        doc.add_paragraph(f"- {c.description} ({c.duration})")

    if medical_card.diagnosis:
        doc.add_heading("Diagnosis", level=2)
        doc.add_paragraph(medical_card.diagnosis.preliminary)

    doc.add_heading("Prescriptions", level=2)
    for p in medical_card.prescriptions:
        doc.add_paragraph(f"- {p.name} ({p.dosage}, {p.duration})")

    if medical_card.doctor_notes:
        doc.add_heading("Doctor notes", level=2)
        doc.add_paragraph(medical_card.doctor_notes)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()


def generate_pdf(medical_card: MedicalCard) -> bytes:
    buffer = io.BytesIO()

    # === REGISTER UNICODE FONT ===
    pdfmetrics.registerFont(
        TTFont("DejaVu", "fonts/DejaVuSans.ttf")
    )

    doc = SimpleDocTemplate(buffer, pagesize=A4)

    # === CUSTOM STYLES WITH UNICODE FONT ===
    title_style = ParagraphStyle(
        name="TitleStyle",
        fontName="DejaVu",
        fontSize=16,
        spaceAfter=14,
        alignment=1  # center
    )

    heading_style = ParagraphStyle(
        name="HeadingStyle",
        fontName="DejaVu",
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6
    )

    normal_style = ParagraphStyle(
        name="NormalStyle",
        fontName="DejaVu",
        fontSize=10,
        spaceAfter=4
    )

    content = []

    # === TITLE ===
    content.append(Paragraph("Medical Record", title_style))

    # === PATIENT INFO ===
    content.append(Paragraph(f"Patient: {medical_card.patient.full_name or '—'}", normal_style))
    content.append(Paragraph(f"Age: {medical_card.patient.age or '—'}", normal_style))
    content.append(Paragraph(f"Gender: {medical_card.patient.gender or '—'}", normal_style))

    # === COMPLAINTS ===
    content.append(Paragraph("Complaints", heading_style))
    for c in medical_card.complaints:
        content.append(
            Paragraph(f"- {c.description} ({c.duration})", normal_style)
        )

    # === DIAGNOSIS ===
    if medical_card.diagnosis:
        content.append(Paragraph("Diagnosis", heading_style))
        content.append(
            Paragraph(medical_card.diagnosis.preliminary, normal_style)
        )

    # === PRESCRIPTIONS ===
    content.append(Paragraph("Prescriptions", heading_style))
    for p in medical_card.prescriptions:
        content.append(
            Paragraph(f"- {p.name} ({p.dosage}, {p.duration})", normal_style)
        )

    # === DOCTOR NOTES ===
    if medical_card.doctor_notes:
        content.append(Paragraph("Doctor notes", heading_style))
        content.append(
            Paragraph(medical_card.doctor_notes, normal_style)
        )

    doc.build(content)
    buffer.seek(0)

    return buffer.read()
